import zipfile
import os
import shutil
from docx import Document

docx_path = r"c:\Users\TranHuy_WIN\Desktop\booking-vietnam\picture.docx"
output_dir = r"c:\Users\TranHuy_WIN\Desktop\booking-vietnam\docx_extracted"

# Clean and create output dir
if os.path.exists(output_dir):
    shutil.rmtree(output_dir)
os.makedirs(output_dir)

# Extract images from docx (which is a zip)
with zipfile.ZipFile(docx_path, 'r') as z:
    for name in z.namelist():
        if name.startswith('word/media/'):
            filename = os.path.basename(name)
            if filename:
                out_path = os.path.join(output_dir, filename)
                with z.open(name) as src, open(out_path, 'wb') as dst:
                    shutil.copyfileobj(src, dst)
                print(f"Extracted: {filename} ({os.path.getsize(out_path)} bytes)")

# Also extract text content to understand structure
print("\n\n=== DOCUMENT TEXT CONTENT ===\n")
doc = Document(docx_path)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[Para {i}]: {para.text.strip()}")

# Check tables
print("\n\n=== TABLES ===\n")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- Table {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  Row {r_idx}, Col {c_idx}: {text[:200]}")

print(f"\n\nTotal images extracted: {len(os.listdir(output_dir))}")
print(f"Images: {sorted(os.listdir(output_dir))}")
